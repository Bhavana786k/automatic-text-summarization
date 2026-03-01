import streamlit as st
import cohere
import PyPDF2
import base64
import io
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# ================= CONFIG =================
st.set_page_config(page_title="Automatic Text Summarization Tool", layout="wide")
st.title("📄 Automatic Text Summarization + Q&A Tool")

co = cohere.Client(st.secrets["cohere"]["api_key"])
MAX_WORD_LIMIT = 15000

# ================= SESSION INIT =================
def init_state():
    defaults = {
        "text": None,
        "output": None,
        "input_mode": None,
        "file_bytes": None,
        "file_type": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ================= RESET =================
def reset_app():
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

# ================= DISPLAY =================
def display_pdf(file_bytes):
    base64_pdf = base64.b64encode(file_bytes).decode("utf-8")
    pdf_html = f"""
        <iframe src="data:application/pdf;base64,{base64_pdf}"
        width="100%" height="900px" style="border:none;"></iframe>
    """
    st.markdown(pdf_html, unsafe_allow_html=True)

def display_text_preview(text):
    st.text_area("File Preview", text, height=900, disabled=True)

# ================= EXTRACTION =================
def extract_pdf(file_bytes):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()

def extract_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs]).strip()

def extract_txt(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore").strip()

# ================= COHERE =================
def cohere_chat(prompt, temperature=0.4, max_tokens=800):
    try:
        response = co.chat(
            model="command-xlarge-nightly",
            message=prompt,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return response.text.strip()
    except Exception as e:
        return f"API Error: {str(e)}"

# ================= SUMMARIZE =================
def summarize_document(text):
    prompt = f"Summarize clearly and concisely:\n\n{text}"
    return cohere_chat(prompt)

# ================= Q&A =================
def generate_qa(text, count):
    prompt = f"""
Generate {count} high-quality questions and answers.

Format strictly:
Q1:
Answer...
Q2:
Answer...

Text:
{text}
"""
    return cohere_chat(prompt, max_tokens=1200)

# ================= DOCX FORMAT =================
def generate_docx(content, mode):
    doc = Document()

    title = doc.add_heading(
        "Document Summary" if mode == "📄 Summarize" else "Questions & Answers",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if mode == "📄 Summarize":
        doc.add_paragraph(content)

    else:
        lines = content.split("\n")
        question = None
        answer_lines = []

        for line in lines:
            line = line.strip()

            if line.startswith("Q"):
                if question:
                    doc.add_heading(question, level=2)
                    doc.add_paragraph(" ".join(answer_lines))
                question = line
                answer_lines = []
            elif line:
                answer_lines.append(line)

        if question:
            doc.add_heading(question, level=2)
            doc.add_paragraph(" ".join(answer_lines))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ================= PDF FORMAT =================
def generate_pdf(content, mode):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)

    styles = getSampleStyleSheet()
    story = []

    heading_style = styles["Heading1"]
    normal_style = styles["Normal"]

    story.append(Paragraph(
        "Document Summary" if mode == "📄 Summarize" else "Questions & Answers",
        heading_style
    ))
    story.append(Spacer(1, 0.5 * inch))

    if mode == "📄 Summarize":
        story.append(Paragraph(content.replace("&", "&amp;"), normal_style))

    else:
        lines = content.split("\n")
        question = None
        answer_lines = []

        for line in lines:
            line = line.strip()

            if line.startswith("Q"):
                if question:
                    story.append(Paragraph(f"<b>{question}</b>", normal_style))
                    story.append(Spacer(1, 0.2 * inch))
                    story.append(Paragraph(" ".join(answer_lines), normal_style))
                    story.append(Spacer(1, 0.4 * inch))
                question = line
                answer_lines = []
            elif line:
                answer_lines.append(line)

        if question:
            story.append(Paragraph(f"<b>{question}</b>", normal_style))
            story.append(Spacer(1, 0.2 * inch))
            story.append(Paragraph(" ".join(answer_lines), normal_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

# ================= CSV =================
def generate_csv(content):
    rows = []
    lines = content.split("\n")
    question = None
    answer_lines = []

    for line in lines:
        line = line.strip()
        if line.startswith("Q"):
            if question:
                rows.append((question, " ".join(answer_lines)))
            question = line
            answer_lines = []
        elif line:
            answer_lines.append(line)

    if question:
        rows.append((question, " ".join(answer_lines)))

    csv_content = "Question,Answer\n"
    for q, a in rows:
        csv_content += f'"{q}","{a}"\n'

    return csv_content.encode("utf-8")

# ================= INPUT =================
uploaded_file = st.file_uploader(
    "Upload a file (PDF, DOCX, TXT)",
    type=["pdf", "docx", "txt"],
)

paste_text = st.text_area("Or paste your text here:", height=200)

word_count = len(paste_text.split()) if paste_text else 0
st.caption(f"Word Count: {word_count} / {MAX_WORD_LIMIT}")

if word_count > MAX_WORD_LIMIT:
    st.error("Word limit exceeded.")

submit_text = st.button("Submit Text", disabled=word_count > MAX_WORD_LIMIT)

if st.button("Reset Application"):
    reset_app()

if uploaded_file and paste_text:
    st.error("Provide either file OR pasted text.")
    st.stop()

# ================= PROCESS INPUT =================
if uploaded_file and not paste_text:
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.name.split(".")[-1].lower()

    st.session_state.file_bytes = file_bytes
    st.session_state.file_type = file_type
    st.session_state.input_mode = "file"

    if file_type == "pdf":
        st.session_state.text = extract_pdf(file_bytes)
    elif file_type == "docx":
        st.session_state.text = extract_docx(file_bytes)
    elif file_type == "txt":
        st.session_state.text = extract_txt(file_bytes)

if paste_text and not uploaded_file and submit_text:
    st.session_state.text = paste_text.strip()
    st.session_state.input_mode = "paste"

# ================= MAIN =================
if st.session_state.text:

    if st.session_state.input_mode == "file":
        col1, col2 = st.columns([1.4, 1])
        with col1:
            if st.session_state.file_type == "pdf":
                display_pdf(st.session_state.file_bytes)
            else:
                display_text_preview(st.session_state.text)
        ai_area = col2
    else:
        ai_area = st.container()

    with ai_area:
        st.markdown("## 🤖 AI Assistant")
        mode = st.radio("Mode:", ["📄 Summarize", "❓ Q&A"])

        if mode == "📄 Summarize":
            if st.button("Generate Summary"):
                with st.spinner("Processing..."):
                    st.session_state.output = summarize_document(st.session_state.text)

        else:
            count = st.slider("Number of Questions", 1, 10, 3)
            if st.button("Generate Q&A"):
                with st.spinner("Generating..."):
                    st.session_state.output = generate_qa(st.session_state.text, count)

        if st.session_state.output:
            st.markdown("---")
            st.subheader("📌 Result")

            if mode == "📄 Summarize":
                st.markdown(st.session_state.output)
            else:
                lines = st.session_state.output.split("\n")
                question = None
                answer_lines = []

                for line in lines:
                    line = line.strip()
                    if line.startswith("Q"):
                        if question:
                            st.markdown(f"### {question}")
                            st.markdown(" ".join(answer_lines))
                            st.markdown("---")
                        question = line
                        answer_lines = []
                    elif line:
                        answer_lines.append(line)

                if question:
                    st.markdown(f"### {question}")
                    st.markdown(" ".join(answer_lines))
                    st.markdown("---")

            formats = ["txt", "docx", "pdf"] if mode == "📄 Summarize" else ["txt", "docx", "pdf", "csv"]
            choice = st.selectbox("Download as:", formats)

            if choice == "txt":
                data = st.session_state.output.encode("utf-8")
            elif choice == "docx":
                data = generate_docx(st.session_state.output, mode)
            elif choice == "pdf":
                data = generate_pdf(st.session_state.output, mode)
            else:
                data = generate_csv(st.session_state.output)

            mime_map = {
                "txt": "text/plain",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "pdf": "application/pdf",
                "csv": "text/csv",
            }

            st.download_button(
                "Download File",
                data=data,
                file_name=f"output.{choice}",
                mime=mime_map[choice],
            )

else:
    st.info("Upload a file OR paste text to begin.")